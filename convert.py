# -*- coding: utf-8 -*-

import sys
import docx 
import importlib
import os
import pandas as pd
import win32com.client
importlib.reload(sys)


from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO

def pdf_to_txt(path):
    fullTxt = []
    
    rsrcmgr = PDFResourceManager()
    
    retstr = StringIO()
    
    codec = 'utf-8'
    
    laparams = LAParams()
    
    device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    
    fp = open(path, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    
    password = ""
    maxpages = 0
    caching = True
    pagenos=set()

    for page in PDFPage.get_pages(fp, 
                                  pagenos, 
                                  maxpages=maxpages,     
                                  password=password,
                                  caching=caching, 
                                  check_extractable=True
    ):
        
        interpreter.process_page(page)

    fullTxt.append(retstr.getvalue().replace("▪", "").\
                                     replace("", "").\
                                     replace("–", "-").\
                                     replace("", "").\
                                     replace("’", "'").\
                                     replace("·", "").\
                                     replace("●", "").\
                                     replace("•", "").\
                                     replace("“", "'").\
                                     replace("”", "'").\
                                     replace("é", "e")
    )
    

    fp.close()
    device.close()
    retstr.close()
    
    return fullTxt


def docx_converter(path):
    
    fullText = []
    # open the docx file
    doc = docx.Document(path)
    
    section = doc.sections[0]
    
    header = section.header
    
    for paragraph in header.paragraphs:
        fullText.append(paragraph.text)
    #
    # read the doc file
    for paragraph in doc.paragraphs:
        
        fullText.append(paragraph.text.replace("▪", "").\
                                           replace("", "").\
                                           replace("–", "-").\
                                           replace("", "").\
                                           replace("’", "'").\
                                           replace("·", "").\
                                           replace("●", "").\
                                           replace("•", "").\
                                           replace("“", "'").\
                                           replace("”", "'").\
                                           replace("é", "e")
                                           )

    return fullText


def doc_to_docx(path):
    
    w = win32com.client.Dispatch('Word.Application')
    
    w.Visible = 0
    w.DisplayAlerts = 0
    
    doc = w.Documents.Open(path)
    
    newpath = os.path.splitext(path)[0] + '.docx'
    
    doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
    
    doc.Close()
    w.Quit()
    os.remove(path)
    
    return newpath




def main():
    
    folder_path = "C:\\Users\\raymond\\Desktop\\Data Resume (no interview)"
    partial_path = "C:\\Users\\raymond\\Desktop\\Data Resume (no interview)\\"
    df = pd.DataFrame(columns = ["Filename", "Content"])
    #
    # read all files in the folder
    for filename in os.listdir(folder_path):
        
        full_path = partial_path + str(filename)
        #
        # identify the file type
        file_extension = os.path.splitext(full_path)[1]
        #
        # if the filename type is pdf
        if file_extension == ".pdf":
            
            output = pdf_to_txt(full_path)
            
            for text in output:
                df = df.append(
                        {
                                "Filename":filename, 
                                "Content":text
                        }, 
                                ignore_index=True
                )
            #
            #seperated by white space
            df = df.append({"Filename":"", "Content":""}, ignore_index=True)
        #
        # the file type is docx
        elif file_extension == ".1docx":
            
            docx_converter(full_path)
            
            output = docx_converter(full_path)
            
            for text in output:
                df = df.append(
                        {
                                "Filename":filename, 
                                "Content":text
                        }, 
                                ignore_index=True
                )
            #
            #seperated by white space
            df = df.append({"Filename":"", "Content":""}, ignore_index=True)
            
        elif file_extension == ".1doc":
            
            new_path = doc_to_docx(full_path)
            
            output = docx_converter(new_path)
            
            for text in output:
                df = df.append(
                        {
                                "Filename":filename, 
                                "Content":text
                        }, 
                                ignore_index=True
                )
            #
            #seperated by white space
            df = df.append({"Filename":"", "Content":""}, ignore_index=True)
            
                
    df.to_csv("C:\\Users\\raymond\\Desktop\\Data Resume (no interview)_pdf_converted.csv", index=False)
    
    return
            
            
if __name__=='__main__':
    main()
                        